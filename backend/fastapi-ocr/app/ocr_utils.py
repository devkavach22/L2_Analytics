# ============================
# OCR & UNIVERSAL FILE PARSER
# ============================

import os
import io
import re
import requests
import fitz  # PyMuPDF
import easyocr
import pandas as pd
import docx2txt
import base64
from datetime import datetime
from PIL import Image
from docx import Document
from pymongo import MongoClient
from dotenv import load_dotenv
from fpdf import FPDF
from bs4 import BeautifulSoup
from typing import Optional, List, Dict, Any

from langdetect import detect, LangDetectException
from deep_translator import GoogleTranslator

from youtube_transcript_api import (
    YouTubeTranscriptApi,
    TranscriptsDisabled,
    NoTranscriptFound,
)

# ----------------------------
# ENV & DB
# ----------------------------
load_dotenv()

MONGO_URL = os.getenv("MONGO_URL")
MONGO_DB_NAME = os.getenv("MONGO_DB_NAME")
OCR_COLLECTION = os.getenv("OCR_COLLECTION", "ocrrecords")

mongo_client = MongoClient(MONGO_URL)
db = mongo_client[MONGO_DB_NAME]
collection = db[OCR_COLLECTION]

# ----------------------------
# OCR INIT
# ----------------------------
READER_LATIN = easyocr.Reader(['en', 'hi'], gpu=False)

# ----------------------------
# OCR HELPERS
# ----------------------------
def ocr_image_bytes(image_bytes: bytes) -> str:
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    import numpy as np
    arr = np.array(img)
    text = READER_LATIN.readtext(arr, detail=0)
    return " ".join(text).strip()

# ----------------------------
# TRANSLATION
# ----------------------------
def detect_and_translate(text: str) -> str:
    if not text or len(text.strip()) < 10:
        return text

    try:
        lang = detect(text[:1000])
        if lang == "en":
            return text

        translator = GoogleTranslator(source='auto', target='en')
        translated = translator.translate(text)

        return (
            f"[ORIGINAL TEXT - {lang}]\n{text}\n\n"
            f"-----------------------------\n"
            f"[TRANSLATED TEXT - EN]\n{translated}"
        )
    except:
        return text

# ----------------------------
# PDF EXTRACTION (FULLY FIXED)
# ----------------------------
def extract_pdf(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    final_pages = []

    for page_index, page in enumerate(doc):
        page_text_parts = []

        # 1️⃣ Extract embedded text
        embedded = page.get_text("text").strip()
        if embedded:
            page_text_parts.append(embedded)

        # 2️⃣ Extract tables (text blocks)
        blocks = page.get_text("blocks")
        for block in blocks:
            if block[6] == 0:
                page_text_parts.append(block[4])

        # 3️⃣ OCR fallback if weak content
        combined = "\n".join(page_text_parts).strip()
        if len(combined) < 200:
            pix = page.get_pixmap(dpi=300)
            ocr_text = ocr_image_bytes(pix.tobytes("png"))
            if ocr_text:
                page_text_parts.append("[OCR PAGE CONTENT]\n" + ocr_text)

        final_pages.append(
            f"\n\n===== PAGE {page_index + 1} =====\n\n" +
            "\n".join(set(page_text_parts))
        )

    return "\n".join(final_pages).strip()

# ----------------------------
# DOCX (TEXT + TABLES)
# ----------------------------
def extract_docx(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    lines = []

    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)

# ----------------------------
# IMAGE PREVIEW (UNCHANGED)
# ----------------------------
def get_file_preview_image(file_bytes: bytes, filename: str) -> Optional[str]:
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext in [".jpg", ".jpeg", ".png", ".bmp", ".webp"]:
            return base64.b64encode(file_bytes).decode()

        if ext == ".pdf":
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pix = doc[0].get_pixmap(dpi=150)
            return base64.b64encode(pix.tobytes("png")).decode()
    except:
        return None

# ----------------------------
# UNIVERSAL EXTRACTOR
# ----------------------------
def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        text = extract_pdf(file_bytes)
    elif ext == ".docx":
        text = extract_docx(file_bytes)
    elif ext in [".jpg", ".jpeg", ".png", ".webp"]:
        text = ocr_image_bytes(file_bytes)
    elif ext in [".txt", ".csv", ".json", ".md", ".log"]:
        text = file_bytes.decode(errors="ignore")
    else:
        text = ocr_image_bytes(file_bytes)

    return detect_and_translate(text)

# ----------------------------
# SAVE TO MONGO
# ----------------------------
def save_record_to_mongo(user_id, filename, text, image=None):
    record = {
        "userId": user_id,
        "filename": filename,
        "extractedText": text,
        "imagePreview": image,
        "createdAt": datetime.utcnow()
    }
    return str(collection.insert_one(record).inserted_id)

# ----------------------------
# MASTER PIPELINE
# ----------------------------
def process_and_save(user_id, filename, file_bytes):
    text = extract_text_from_file(file_bytes, filename)
    image = get_file_preview_image(file_bytes, filename)
    record_id = save_record_to_mongo(user_id, filename, text, image)

    return {
        "success": True,
        "recordId": record_id,
        "text": text,
        "imagePreview": image
    }

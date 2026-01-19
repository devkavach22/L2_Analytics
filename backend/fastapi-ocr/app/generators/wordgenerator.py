# generators/word_generator.py
from docx import Document
import os

def create_word(text_content: str, filepath: str):
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc = Document()
    for line in text_content.splitlines():
        line = line.strip()
        if not line: continue
        if line.endswith(":") or line.isupper():
            doc.add_heading(line.replace(":", ""), level=2)
        else:
            doc.add_paragraph(line)
    doc.save(filepath)
    return filepath
